import folium
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from PIL import Image

from .rosreestr2 import GetArea


def get_map(number_list:list):
    m = folium.Map(location=[55.7558, 37.6173], zoom_start=6, zoom_control=False,
                   control_scale=True)

    m.options.update({'max_width': '100%'})
    m.get_root().html.add_child(
        folium.Element("<style>.leaflet-control-attribution.leaflet-control{display:none;}</style>"))

    all_place_lat = []
    all_place_lng = []
    for number in number_list:
        areas = GetArea(number)
        coordinates = areas.get_coord()
        if coordinates:
            for coordinate in coordinates:
                for addresses in coordinate:
                    points = []
                    for pt in addresses:
                        place_lat = pt[1]
                        place_lng = pt[0]
                        all_place_lat.append(place_lat)
                        all_place_lng.append(place_lng)
                        points.append([place_lat, place_lng])
                    folium.Polygon(points, color='red').add_to(m)

                    center_point_lng = areas.center['x'],
                    center_point_lat = areas.center['y'],
                    folium.CircleMarker(
                        location=[center_point_lat[0], center_point_lng[0]],
                        popup=folium.Popup(f':{number.split(":")[-1]}', show=True),
                        opacity=0,
                    ).add_to(m)

    m.get_root().html.add_child(
        folium.Element("<style>.leaflet-popup-close-button {display: none;}</style>"))
    bounds = [[min(all_place_lat), min(all_place_lng)], [
        max(all_place_lat), max(all_place_lng)]]
    center_lat = (bounds[0][0] + bounds[1][0]) / 2
    center_lng = (bounds[0][1] + bounds[1][1]) / 2
    m.location = [center_lat, center_lng]
    m.fit_bounds(bounds)

    map_html = m._repr_html_()

    # order = get_object_or_404(Order, id=order_id)
    # order.map = map_html
    # order.save()

    return map_html


# Получаем объект карты, для сохранения скриншота к заказу
def get_map_screenshot(number_list:list):
    m = folium.Map(location=[55.7558, 37.6173], zoom_start=6, zoom_control=False,
                   control_scale=True)

    m.options.update({'max_width': '100%'})
    m.get_root().html.add_child(
        folium.Element("<style>.leaflet-control-attribution.leaflet-control{display:none;}</style>"))

    all_place_lat = []
    all_place_lng = []
    for number in number_list:
        areas = GetArea(number)
        coordinates = areas.get_coord()
        if coordinates:
            for coordinate in coordinates:
                for addresses in coordinate:
                    points = []
                    for pt in addresses:
                        place_lat = pt[1]
                        place_lng = pt[0]
                        all_place_lat.append(place_lat)
                        all_place_lng.append(place_lng)
                        points.append([place_lat, place_lng])
                    folium.Polygon(points, color='red').add_to(m)

                    center_point_lng = areas.center['x'],
                    center_point_lat = areas.center['y'],
                    folium.CircleMarker(
                        location=[center_point_lat[0], center_point_lng[0]],
                        popup=folium.Popup(f':{number.split(":")[-1]}', show=True),
                        opacity=0,
                    ).add_to(m)

    m.get_root().html.add_child(
        folium.Element("<style>.leaflet-popup-close-button {display: none;}</style>"))
    bounds = [[min(all_place_lat), min(all_place_lng)], [
        max(all_place_lat), max(all_place_lng)]]
    center_lat = (bounds[0][0] + bounds[1][0]) / 2
    center_lng = (bounds[0][1] + bounds[1][1]) / 2
    m.location = [center_lat, center_lng]
    m.fit_bounds(bounds)

    return m



# Выгрузка DOCX
# Замена заполнителей значениями в абзаце.
def replace_placeholders(paragraph:str, placeholders:dict):
    for placeholder, value in placeholders.items():
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    if placeholder == '_обзорная_схема':
                        run.text = ""
                        width, height = Image.open(value).size
                        run.add_picture(value, width=Inches(width / 192), height=Inches(height / 192))
                    else:
                        run.text = run.text.replace(placeholder, value)


# Замена заполнителей значениями в таблице.
def replace_placeholders_in_table(table:str, placeholders:dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders(paragraph, placeholders)


# Замена заполнителей значениями в футере документа.
def replace_placeholders_in_footer(document, placeholders:dict):
    sections = document.sections
    for section in sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_placeholders(paragraph, placeholders)


# Замена заполнителей значениями во всех абзацах и таблицах документа
def replace_placeholders_in_document(document, placeholders:dict):
    for paragraph in document.paragraphs:
        replace_placeholders(paragraph, placeholders)

    for table in document.tables:
        replace_placeholders_in_table(table, placeholders)

    replace_placeholders_in_footer(document, placeholders)

    return document


# Генерация нового документа с заменой заполнителей значениями.
def generate_docx(document_path, placeholders:dict):
    document = Document(document_path)
    replace_placeholders_in_document(document, placeholders)
    return document


def add_table(document, coordinates_dict:dict):
    for paragraph in document.paragraphs:
        if '_таблица_координат' in paragraph.text:
            for key in coordinates_dict:
                document.add_paragraph(f'Координаты углов участка {key}', style='Normal')
                table = document.add_table(rows=len(coordinates_dict[key]) + 1, cols=3)

                # Добавляем границы таблицы
                table.style = 'Table Grid'

                # Выравниваем таблицу по центру
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Добавляем строку с названиями столбцов и выравниваем их по центру
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Номер точки'
                hdr_cells[1].text = 'Координата Х'
                hdr_cells[2].text = 'Координата У'
                for cell in hdr_cells:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Добавляем данные в ячейки таблицы и выравниваем их по центру
                for i, coord in enumerate(coordinates_dict[key]):
                    row_cells = table.rows[i + 1].cells
                    row_cells[0].text = str(i + 1)
                    row_cells[1].text = str(coord[0])
                    row_cells[2].text = str(coord[1])
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Вставляем таблицу после абзаца, содержащего "_таблица координат".
                document.add_paragraph('', style='Normal')

                # Удаляем абзац с заполнителем таблицы
                paragraph._element.clear()
            break